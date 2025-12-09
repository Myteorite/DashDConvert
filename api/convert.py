# api/convert.py
from fastapi import FastAPI, File, UploadFile, Form
from fastapi.responses import StreamingResponse
import shutil
import tempfile
import os
from pathlib import Path
from io import BytesIO

# converters
from docx import Document as DocxDocument
from PyPDF2 import PdfReader, PdfWriter
from fpdf import FPDF

app = FastAPI()

def docx_to_text(path: str) -> str:
    doc = DocxDocument(path)
    texts = []
    for p in doc.paragraphs:
        texts.append(p.text)
    return "\n".join(texts)

def pdf_to_text(path: str) -> str:
    reader = PdfReader(path)
    texts = []
    for page in reader.pages:
        try:
            texts.append(page.extract_text() or "")
        except Exception:
            texts.append("")
    return "\n".join(texts)

def text_to_pdf_bytes(text: str) -> bytes:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    # split lines to avoid overflow
    for line in text.splitlines():
        pdf.multi_cell(0, 6, line)
    b = BytesIO()
    pdf.output(b)
    return b.getvalue()

def text_to_docx_bytes(text: str) -> bytes:
    doc = DocxDocument()
    for line in text.splitlines():
        doc.add_paragraph(line)
    b = BytesIO()
    doc.save(b)
    b.seek(0)
    return b.read()

@app.post("/api/convert")
async def convert(file: UploadFile = File(...), to_format: str = Form(...)):
    # to_format: 'txt', 'pdf', 'docx'
    # small safety: accept only a few content types
    filename = Path(file.filename).name
    suffix = Path(filename).suffix.lower()
    with tempfile.TemporaryDirectory() as tmp:
        in_path = os.path.join(tmp, filename)
        # save uploaded file
        with open(in_path, "wb") as f:
            shutil.copyfileobj(file.file, f)

        # extract text depending on input
        text = ""
        if suffix in [".docx"]:
            text = docx_to_text(in_path)
        elif suffix in [".pdf"]:
            text = pdf_to_text(in_path)
        elif suffix in [".txt", ".md", ".csv"]:
            with open(in_path, "r", encoding="utf-8", errors="ignore") as fr:
                text = fr.read()
        else:
            return {"error": "Unsupported input format. Use .pdf, .docx, .txt"}

        # produce output bytes & filename
        out_bytes = None
        out_name = None
        if to_format == "txt":
            out_bytes = text.encode("utf-8")
            out_name = Path(filename).stem + ".txt"
            media_type = "text/plain"
        elif to_format == "pdf":
            out_bytes = text_to_pdf_bytes(text)
            out_name = Path(filename).stem + ".pdf"
            media_type = "application/pdf"
        elif to_format == "docx":
            out_bytes = text_to_docx_bytes(text)
            out_name = Path(filename).stem + ".docx"
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            return {"error": "Unsupported target format. Use 'txt', 'pdf', or 'docx'."}

        # return as streaming response for download
        return StreamingResponse(BytesIO(out_bytes), media_type=media_type,
                                 headers={"Content-Disposition": f'attachment; filename="{out_name}"'})
